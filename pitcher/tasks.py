# -*- coding: utf-8 -*-
"""
Created on Mon Nov 17 20:03:38 2014

@author: Administrator
"""

#%%
from __future__ import division
import time
from datetime import datetime, timedelta
from ticketpitcher import pitcher
from pitcher.models import *
from pandas import DataFrame
from dateutil import parser
from django_pandas.io import read_frame

import tempfile
from docx import Document
from docx.shared import Inches
from BeautifulSoup import BeautifulSoup
from time import sleep
from django.conf import settings
from mail import send_mail


class PitchTask():
    '''
    抢票任务
    '''

    def __init__(self, taskName):
        #%% 设置全局变量
        systemConfig = SystemConfig.objects.all()[0]
        self.normalWaitingSecond = systemConfig.normalWaitingSecond
        self.errorWaitingSecond = systemConfig.errorWaitingSecond
        self.maxLoginError = systemConfig.maxLoginError
        self.maxException = systemConfig.maxException
        self.timeToStop = systemConfig.timeToStop
        # 读取信息信息
        task = Task.objects.get(taskName=taskName)
        self.task = task
        self.username = task.username
        self.password = task.password
        self.working = task.working
        date = datetime.now() + timedelta(systemConfig.preceding)
        self.day = datetime.strftime(date, "%Y-%m-%d")
        self.pitchConfig = self.getPitchConfig()


    def getPitchConfig(self):
        '''
        获取要抢票的配置
        返回DataFrame数据集,数据已按优先级排好顺序
        注意：排序规则在数据模型PitchConfig中的ordering定义的
        '''
        data = []
        for pitchConfig in PitchConfig.objects.filter(need__gt=0, task=self.task):
            flight = pitchConfig.flight
            data.append([flight.flightCode, pitchConfig.need])
        return DataFrame(data, columns=['flightCode', 'need'])

    def writeSystemLog(self, msg):
        '''
        写入系统日志
        '''
        # 将信息写到数据库日志中
        log = SystemLog()
        log.logMsg = msg
        log.task = self.task
        log.save()


    def isLogin(self):
        '''
        检查当前是否已经登录
        '''
        return pitcher.isLogin()


    def login(self, username, password):
        '''
        登录
        '''
        return pitcher.login(username, password)


    def getTicketInfo(self, day):
        '''
        按天获取船票的信息
        day 日期格式为'yyyy-mm-dd'
        '''
        return pitcher.getTicketInfo(day)


    def orderTicket(self, dailyFlightId, n):
        '''
        订票
        '''
        return pitcher.orderTicket(dailyFlightId, n)


    def getTicketRemain(self, ticketInfo, flightCode):
        '''
        根据提供的船票信息数据集及航班的编码获得余票数
        ticketInfo 船票信息
        flightCode 航班的编码
        '''
        #df = ticketInfo[ticketInfo[u'航班号'] == flightCode]
        df = ticketInfo[ ticketInfo[u'航班号'].apply(lambda x : x.startswith(flightCode)) ]
        if len(df.index) > 0:
            return int(df.irow(0)[u'余票'])
        else:
            return 0


    def getDailyFlightId(self, ticketInfo, flightCode):
        '''
        根据提供的船票信息数据集及航班的编码获得dailyFlightId
        ticketInfo 船票信息
        flightCode 航班的编码
        '''
        #df = ticketInfo[ticketInfo[u'航班号'] == flightCode]
        df = ticketInfo[ ticketInfo[u'航班号'].apply(lambda x : x.startswith(flightCode)) ]
        if len(df.index) > 0:
            return df.irow(0)[u'航班ID'][1:-1]
        else:
            return None


    def pitchItem(self, ticketInfo, flightCode, need):
        '''
        进行一次抢票动作
        ticketInfo  最近读取的船票信息
        flightCode  航班代码
        n           需要票的数量
        成功返回实际抢票数量，失败返回0
        '''
        ItemMessage = u'尝试抢票(航班号:%s,需票数:%d):' % (flightCode, need)
        self.writeSystemLog(ItemMessage + u'开始...')
        if need == 0:
            self.writeSystemLog(ItemMessage + u'出错: 需票数为0')
            return 0
        # 读取票数剩余
        remain = self.getTicketRemain(ticketInfo, flightCode)
        # 如果该航班没有票跳过
        if remain == 0:
            self.writeSystemLog(ItemMessage + u'已无余票.')
            return 0
        # 获取dailyFlightId
        dailyFlightId = self.getDailyFlightId(ticketInfo, flightCode)
        if dailyFlightId == None:
            self.writeSystemLog(ItemMessage + u'无法获取航班Id.')
            return 0
        # 如果需票数小于实际剩余按实际票数抢
        if need > remain:
            self.writeSystemLog(ItemMessage + u'余票不足，按实际票数抢.')
            n = remain
        else:
            n = need
        #writeSystemLog(ItemMessage + u'dailyFlightId:%s' % dailyFlightId)
        orderResult = self.orderTicket(dailyFlightId, n)
        if orderResult:
            return n
        else:
            return 0


    def savePitchLog(self, ticketInfo, flightCode, need, pitchResult):
        '''
        保存抢票记录信息
        ticketInfo   最新的船票信息数据集
        flightCode   航班号
        need         需票数
        pitchResult  实际抢票数量
        '''
        # 记录本次抢票结果
        pitchLog = PitchLog()
        pitchLog.task = self.task
        pitchLog.flightCode = flightCode
        pitchLog.need = need
        pitchLog.pitchCount = pitchResult
        #df = ticketInfo[ticketInfo[u'航班号'] == flightCode]
        df = ticketInfo[ ticketInfo[u'航班号'].apply(lambda x : x.startswith(flightCode)) ]
        if len(df.index) > 0:
            row = df.irow(0)
            pitchLog.flightId = row[u'航班ID']
            pitchLog.departure = row[u'出发码头']
            pitchLog.arrival = row[u'抵达码头']
            pitchLog.departureTime = row[u'开航时间']
            pitchLog.ticketCount = row[u'余票']
            # 保存抢票日志
            pitchLog.save()
        else:
            ItemMessage = u'尝试抢票(航班号:%s,需票数:%d):' % (flightCode, need)
            self.writeSystemLog(ItemMessage + u'无法获取余票及航班相关信息.')


    def pitchLoop(self):
        '''
        刷票主循环过程
        返回True表示抢票动作没有执行，但是登录信息丢失，需要登录后继续调用该过程
        返回False表示抢票动作已执行，或者已经超过限定的抢票时间，程序应该终止
        '''
        while self.isLogin():
            ticketInfo = self.getTicketInfo(self.day)
            if len(ticketInfo.index):
                ### 已经放票了,根据配置数据开始顺序进行抢票 ###
                self.writeSystemLog(u'系统已经放票，准备开始抢票...')
                resultList = []
                for i in range(len(self.pitchConfig.index)):
                    config = self.pitchConfig.irow(i)
                    # 读取当前记录要抢的航班号和需票信息
                    flightCode = config['flightCode']
                    need = config['need']
                    ItemMessage = u'尝试抢票(航班号:%s,需票数:%d):' % (flightCode, need)
                    # 调用抢票接口
                    beginTime = datetime.now()
                    pitchResult = self.pitchItem(ticketInfo, flightCode, need)
                    endTime = datetime.now()
                    # 将抢票结果反映在系统日志中
                    if pitchResult > 0:
                        timeConsumed = endTime - beginTime
                        # 抢票成功
                        self.writeSystemLog(ItemMessage + u'抢票成功!(耗时%.2f秒)' % timeConsumed.total_seconds())
                    else:
                        # 抢票失败
                        self.writeSystemLog(ItemMessage + u'抢票执行失败，将跳过此项.')
                    # 将结果记录列表抢票结束后统一保存（加快速度）
                    resultList.append({'flightCode': flightCode, 'need': need, 'pitchResult': pitchResult})

                ### 保存抢票结果信息 ###
                # 重新读取余票信息
                ticketInfo = self.getTicketInfo(self.day)
                # 保存抢票结果记录信息
                for result in resultList:
                    self.savePitchLog(ticketInfo, result['flightCode'], result['need'], result['pitchResult'])

                ### 执行过抢票程序任务完成 ###
                self.writeSystemLog(u'已完成刷票动作，将停止执行.')
                return False  #返回False表示程序应该终止

            # 刷新完毕等待一段时间
            self.writeSystemLog(u'完成一次刷新，将等待%d秒...' % self.normalWaitingSecond)
            time.sleep(self.normalWaitingSecond)
            # 检查当前时间是否超过中午12点，如果超过停止刷票
            if datetime.strftime(datetime.now(), "%H:%M") >= self.timeToStop:
                self.writeSystemLog(u'程序已执行到执行时间(%s),将停止执行.' % self.timeToStop)
                return False  #返回False表示程序应该终止
        # 程序到这里表示登录信息已经丢失
        self.writeSystemLog(u'登录信息丢失将尝试重新登录.')
        return True

    def exportPayInfo(self):
        """ 导出出付费信息 """
        reserveInfo = pitcher.getReserveInfo()
        c1 = reserveInfo[u'状态'] == u'未确认'
        c2 = reserveInfo[u'最后确认时间'].apply(parser.parse) > datetime.now()
        reserveInfo = reserveInfo[c1 & c2]
        document = Document()
        for i, row in reserveInfo.iterrows():
            document.add_heading(u'票项%d' % (i + 1), level=1)
            document.add_paragraph(text=u'航线: ' + row[u'航线'])
            document.add_paragraph(text=u'航班时间: ' + row[u'航班时间'])
            document.add_paragraph(text=u'人数: ' + row[u'人数'])
            document.add_paragraph(text=u'金额: ' + row[u'金额'])
            document.add_paragraph(text=u'最后确认时间: ' + row[u'最后确认时间'])
            filename = tempfile.mktemp(suffix='.jpg',prefix='tmp_')
            with open(filename, 'wb') as f:
                orderNumber = pitcher.getOrderNumber(row[u'预订ID'])
                qrcode = pitcher.getWeixinPayQrcode(orderNumber)
                f.write(qrcode)
            document.add_picture(filename, width=Inches(1))
            time.sleep(self.normalWaitingSecond)
        filename = tempfile.mktemp(suffix='.docx',prefix='tmp_')
        document.save(filename)

        # 发送邮件
        title = '%s[%s]' % (self.username, datetime.strftime(datetime.now(), "%Y-%m-%d"))
        send_mail(settings.MAIL_LIST, title, u'见附件', [filename])

    def run(self):
        '''
        抢票程序主过程
        '''
        self.writeSystemLog(u'程序开始执行...')
        # 判断刷票开关是否开启，如果没有开启则退出
        if not self.working:
            self.writeSystemLog(u'刷票任务配置未启动，程序将退出.')
            self.writeSystemLog(u'程序执行结束.')
            return

        # 开始刷票
        loginErrorCount = 0
        ExceptionCount = 0
        pitchLoopResult = True
        while pitchLoopResult:
            try:
                loginResult = self.login(self.username, self.password)
                if loginResult:
                    # 如果登录成功清空原来登录失败的记录
                    loginErrorCount = 0
                    self.writeSystemLog(u'登录成功!')
                    pitchLoopResult = self.pitchLoop()
                else:
                    loginErrorCount += 1
                    if loginErrorCount > self.maxLoginError:
                        self.writeSystemLog(u'登录失败超过%d次，程序退出!' % self.maxLoginError)
                        return
                    else:
                        self.writeSystemLog(u'登录失败!')
                        self.writeSystemLog(u'等待%s秒... ...' % self.errorWaitingSecond)
                        time.sleep(self.errorWaitingSecond)
            except Exception as e:
                self.writeSystemLog(u'出现异常:' + unicode(e))
                ExceptionCount += 1
                if ExceptionCount > self.maxException:
                    self.writeSystemLog(u'抛出异常超过%d次，程序将退出!' % self.maxException)
                    break
                self.writeSystemLog(u'等待%s秒... ...' % self.errorWaitingSecond)
                time.sleep(self.errorWaitingSecond)


        # 开始导出支付清单
        for i in range(5):
            try:
                self.exportPayInfo()
                self.writeSystemLog(u'导出支付清单成功!')
                break
            except Exception:
                loginResult = self.login(self.username, self.password)
        else:
            self.writeSystemLog(u'导出支付清单失败!')


        self.writeSystemLog(u'程序执行结束!')


def getFinalTime(x):
    '''
    获取 最终确认时间，即开船前1天的下午2点
    x 开航时间 用文本存储的日期+时间格式（即直接从reserveInfo中读取的'航班时间'
    '''
    beginTime_1 = parser.parse(x) - timedelta(1)
    lastDay = datetime.strftime(beginTime_1, "%Y-%m-%d")
    return parser.parse(lastDay + ' 14:00')


class RefreshTask():
    '''
    更新任务
    '''

    def __init__(self, taskName):
        '''
        构造函数
        '''
        #%% 设置全局变量
        systemConfig = SystemConfig.objects.all()[0]
        self.normalWaitingSecond = systemConfig.normalWaitingSecond
        self.errorWaitingSecond = systemConfig.errorWaitingSecond
        self.maxLoginError = systemConfig.maxLoginError
        self.maxException = systemConfig.maxException
        self.timeToStop = systemConfig.timeToStop
        # 导入任务信息
        task = Task.objects.get(taskName=taskName)
        self.task = task
        self.username = task.username
        self.password = task.password
        self.working = task.working

    def writeSystemLog(self, msg):
        '''
        写入系统日志
        '''
        # 将信息写到数据库日志中
        log = SystemLog()
        log.logMsg = msg
        log.task = self.task
        log.save()

    def isLogin(self):
        '''
        检查当前是否已经登录
        '''
        return pitcher.isLogin()


    def login(self, username, password):
        '''
        登录
        '''
        return pitcher.login(username, password)


    def getTicketInfo(self, day):
        '''
        按天获取船票的信息
        day 日期格式为'yyyy-mm-dd'
        '''
        return pitcher.getTicketInfo(day)


    def orderTicket(self, dailyFlightId, n):
        '''
        订票
        '''
        return pitcher.orderTicket(dailyFlightId, n)


    def refreshReserve(self, reverseId, reserveInfo=None):
        '''
        通过取消预订，再重新预订的方法刷新预订的最后确认时间
        reverseId   要刷新的预订记录的ID
        reserveInfo 当前用户的预订信息（通过调用getReserveInfo获得），如果为空将重新请求
        '''
        if reserveInfo == None:
            reserveInfo = pitcher.getReserveInfo()

        # 读取当前预订数据
        c1 = reserveInfo[u'预订ID'] == str(reverseId)
        data = reserveInfo[c1]
        if len(data) != 1:
            return False

        # 通过预订数据获取航班ID
        beginDay = data.irow(0)[u'开航日期']
        beginTime = data.irow(0)[u'开航时间']
        departure = data.irow(0)[u'出发码头']
        arrival = data.irow(0)[u'抵达码头']
        cnt = int(data.irow(0)[u'人数'])
        self.writeSystemLog(u'出发:%s,抵达:%s,开航时间:%s %s,人数:%s' % (departure, arrival, beginDay, beginTime, cnt))
        dailyFlightId = pitcher.getDailyFlightId(beginDay, beginTime, departure, arrival)
        if dailyFlightId == None:
            self.writeSystemLog(u'获取航班ID失败')
            return False
        self.writeSystemLog(u'dailyFlightId=%d' % dailyFlightId)

        # 取消预订
        if pitcher.cancelReserve(reverseId) == False:
            self.writeSystemLog(u'取消预订失败')
            return False
        self.writeSystemLog(u'预订已取消,将尝试回订')

        # 重新预订
        error = 0
        while pitcher.orderTicket(dailyFlightId, cnt) == False:
            error += 1
            if error > 3:
                self.writeSystemLog(u'回订出错超过3次，将跳过此项!!!')
                # 1月22系统改规则了，退票后两个小时才会放出来所以这里无论做多少次都不会成功的
                # 将数据放入重做日志中后续慢慢刷
                redo = RefreshRedo(beginDay=beginDay, beginTime=beginTime, departure=departure,
                                   arrival=arrival, cnt=cnt, state='redoing')
                redo.save()
                self.writeSystemLog(u'数据已保存到重做日志.')
                return False
            self.writeSystemLog(u'回订出错将重试...dailyFlightId=%d' % dailyFlightId)
            time.sleep(2)
            # 重新读取dailyFlightId，如果出错是由dailyFlightId造成无论重做多少次都是没有用的
            dailyFlightId = pitcher.getDailyFlightId(beginDay, beginTime, departure, arrival)

        self.writeSystemLog(u'回订成功!')
        return True


    def runRefresh(self):
        '''
        更新票项
        1月22系统改规则了，退票后两个小时才会放出
        该过程的实际功能变为把所有票项退掉并记录重做日志
        后续由runRedo再根据重做日志将票抢回
        '''
        self.writeSystemLog(u'程序开始执行...')
        # 判断刷票开关是否开启，如果没有开启则退出
        if not self.working:
            self.writeSystemLog(u'更新任务配置未启动，程序将退出.')
            self.writeSystemLog(u'程序执行结束.')
            return

        # 尝试进程登录
        loginResult = self.login(self.username, self.password)
        if loginResult:
            self.writeSystemLog(u'登录成功!')
        else:
            self.writeSystemLog(u'登录失败，程序将退出.')
            self.writeSystemLog(u'程序执行结束.')
            return

        # 获取当前用户的预订信息
        self.writeSystemLog(u'尝试获取当前用户预订信息...')
        reserveInfo = pitcher.getReserveInfo()
        reserveInfo[u'最后确认时间(最终)'] = reserveInfo[u'航班时间'].apply(getFinalTime)
        now = datetime.now()
        day = datetime.strftime(now, "%Y-%m-%d")
        # 条件1：最后确认时间要 大于 系统时间
        c1 = reserveInfo[u'最后确认时间'].apply(lambda x: parser.parse(x)) > now
        # 条件2：最后确认时间 不等于 最后确认时间（最终）
        c2 = reserveInfo[u'最后确认时间(最终)'] != reserveInfo[u'最后确认时间']
        # 条件3：不是今天订的票（或者今天刚刷新过）
        c3 = reserveInfo[u'预约时间'].apply(lambda x: datetime.strftime(parser.parse(x), "%Y-%m-%d")) != day
        # 获取需要更新的预订
        reserveInfo = reserveInfo[c1 & c2 & c3]
        self.writeSystemLog(u'预订信息获取完毕.')

        # 检查是否有需要更新的预订
        if len(reserveInfo) == 0:
            self.writeSystemLog(u'没有需要更新的预订，程序将退出.')
            self.writeSystemLog(u'程序执行结束.')
            return
        self.writeSystemLog(u'有%d项预订需要更新...' % len(reserveInfo))

        # 开始循环的更新预订票项
        error = 0
        try:
            # 刷新所有的预订的预订时间
            for i in reserveInfo.iterrows():
                # 读取预订相关信息
                row = i[1]
                reserveId = row[u'预订ID']
                #departureTIME = row[u'航班时间']
                #departure = row[u'出发码头']
                #arrival = row[u'抵达码头']
                # 将信息输入到日志以便出错的时候可以手工介入处理
                self.writeSystemLog(u'尝试更新预订信息(reserveId=%s)...' % reserveId)
                #self.writeSystemLog(u'出发:%s,抵达:%s,开航时间:%s' % (departure, arrival, departureTIME))
                if self.refreshReserve(reserveId):
                    self.writeSystemLog(u'更新成功.')
                else:
                    error += 1
                    self.writeSystemLog(u'更新失败,将跳过此项.')
                    # 1月22系统改规则了，退票后两个小时才会放出
                    # 多数票项的刷新都会失败只能依靠出错日志重做了
                    # 所以这里无论失败多少次都不能退出
                    #if error > self.maxException:
                    #    self.writeSystemLog(u'抛出异常超过%d次，程序将退出!' % self.maxException)
                    #    return
                    # 检查连接是否丢失
                    if not self.isLogin():
                        self.writeSystemLog(u'连接丢失，程序将退出.')
                        self.writeSystemLog(u'程序执行结束.')
                        return
                self.writeSystemLog(u'完成一次更新!')
                self.writeSystemLog(u'等待%s秒... ...' % self.normalWaitingSecond)
        except Exception as e:
            self.writeSystemLog(u'刷新中出现异常，程序将退出')
            self.writeSystemLog(u'异常信息为:%s' % unicode(e))
            self.writeSystemLog(u'更新出错%d次(不含异常)' % error)
            self.writeSystemLog(u'程序异常结束.')
            return
        self.writeSystemLog(u'已完成预订信息的刷新，程序将退出.')
        self.writeSystemLog(u'更新出错%d次' % error)
        self.writeSystemLog(u'程序执行结束.')

    def runRedo(self):
        '''
        将更新票项的过程中失败的数据重做直到成功
        '''
        # 查出需要进行重做的数据
        self.writeSystemLog(u'重做过程开始执行')
        # 执行时间变量的初始化
        runStartTime = datetime.now()
        currentTime = datetime.now()
        timeSpend = currentTime - runStartTime
        # 获取总出错项
        redoList = list(RefreshRedo.objects.filter(state=u'redoing'))
        errorCount = len(redoList)
        self.writeSystemLog(u'共有%d次错误需要重做' % errorCount)

        # 不停处理每一个信息项,直到没有错误需要重做,或者执行时间到了
        while errorCount > 0 and timeSpend.total_seconds() / 3600 < 5:
            try:
                redoList = list(RefreshRedo.objects.filter(state=u'redoing'))
                errorCount = len(redoList)
                if errorCount == 0:
                    self.writeSystemLog(u'已无项目需要处理,程序将退出')
                else:
                    self.writeSystemLog(u'开始一次重做，共%d需要处理...' % errorCount)

                # 根据出错列表调用重新回订过程
                successCount = 0
                for redo in redoList:
                    departure = redo.departure
                    arrival = redo.arrival
                    beginDay = redo.beginDay
                    beginTime = redo.beginTime
                    cnt = redo.cnt

                    # 如果当前时间和开航日期0点的时间小于1天，则不再刷新这个票项
                    # 因为正常刷新提前5天前就刷新好了，所以不会有影响
                    # 这样做主要是防止重做票项无限增长，且去订之前日期的票
                    beginDateTime = parser.parse(beginDay)
                    if (beginDateTime - currentTime).days < 1:
                        redo.state = u'failed'
                        redo.save()
                        self.writeSystemLog(
                            u'出发:%s,抵达:%s,开航时间:%s %s,人数:%s' % (departure, arrival, beginDay, beginTime, cnt))
                        self.writeSystemLog(u'该项已经超过可以重做的时间，已标记为失败!!!')
                        continue

                    # 尝试重调订票过程
                    dailyFlightId = pitcher.getDailyFlightId(beginDay, beginTime, departure, arrival)
                    if dailyFlightId:
                        result = pitcher.orderTicket(dailyFlightId, cnt)
                        if result:
                            # 重做成功将记录标识为完成
                            redo.state = u'finished'
                            redo.save()
                            self.writeSystemLog(
                                u'出发:%s,抵达:%s,开航时间:%s %s,人数:%s' % (departure, arrival, beginDay, beginTime, cnt))
                            self.writeSystemLog(u'该项已经重做成功!')
                            successCount += 1

                    # 等待一定时间(避免过于频繁访问服务器)
                    # 由于只有我们知道票出现的大概时间，被别人订走的可能性不大
                    time.sleep(self.normalWaitingSecond)


                # 完成一次重做,打印执行信息
                errorCount -= successCount
                self.writeSystemLog(u'完成一次重做：%d项成功，%d项失败.' % (successCount, errorCount))
                # 更新执行时间
                currentTime = datetime.now()
                timeSpend = currentTime - runStartTime
            except Exception as e:
                self.writeSystemLog(u'发生异常:%s，程序将继续执行' % unicode(e))


            # 检查是否是登录状态，如果登录状态丢失则重新登录
            if not self.isLogin():
                self.writeSystemLog(u'连接丢失，将重新登录...')
                loginErrorCount = 0
                while True:
                    loginResult = self.login(self.username, self.password)
                    if loginResult:
                        # 如果登录成功清空原来登录失败的记录
                        self.writeSystemLog(u'登录成功!')
                        break
                    else:
                        loginErrorCount += 1
                        if loginErrorCount > self.maxLoginError:
                            self.writeSystemLog(u'登录失败超过%d次，程序退出!' % self.maxLoginError)
                            return
                        else:
                            self.writeSystemLog(u'登录失败!')
                            self.writeSystemLog(u'等待%s秒... ...' % self.errorWaitingSecond)
                            time.sleep(self.errorWaitingSecond)

        self.writeSystemLog(u'重做过程执行结束')

    def run(self):
        '''
        更新主过程
        '''
        self.writeSystemLog(u'开始调用刷新票项的过程...')
        self.runRefresh()
        self.writeSystemLog(u'刷新票项的过程已返回...')

        self.writeSystemLog(u'开始重做过程...')
        self.runRedo()
        self.writeSystemLog(u'重做过程已返回...')


class RedoTask():
    '''
    更新任务
    '''

    def __init__(self, taskName):
        '''
        构造函数
        '''
        #%% 设置全局变量
        systemConfig = SystemConfig.objects.all()[0]
        self.normalWaitingSecond = systemConfig.normalWaitingSecond
        self.errorWaitingSecond = systemConfig.errorWaitingSecond
        self.maxLoginError = systemConfig.maxLoginError
        self.maxException = systemConfig.maxException
        self.timeToStop = systemConfig.timeToStop
        # 导入任务信息
        task = Task.objects.get(taskName=taskName)
        self.task = task
        self.username = task.username
        self.password = task.password
        self.working = task.working

    def writeSystemLog(self, msg):
        '''
        写入系统日志
        '''
        # 将信息写到数据库日志中
        log = SystemLog()
        log.logMsg = msg
        log.task = self.task
        log.save()

    def isLogin(self):
        '''
        检查当前是否已经登录
        '''
        return pitcher.isLogin()


    def login(self, username, password):
        '''
        登录
        '''
        return pitcher.login(username, password)


    def getTicketInfo(self, day):
        '''
        按天获取船票的信息
        day 日期格式为'yyyy-mm-dd'
        '''
        return pitcher.getTicketInfo(day)


    def orderTicket(self, dailyFlightId, n):
        '''
        订票
        '''
        return pitcher.orderTicket(dailyFlightId, n)


    def runRedo(self):
        '''
        将更新票项的过程中失败的数据重做直到成功
        '''
        # 查出需要进行重做的数据
        self.writeSystemLog(u'重做过程开始执行')

        # 尝试进程登录
        loginResult = self.login(self.username, self.password)
        if loginResult:
            self.writeSystemLog(u'登录成功!')
        else:
            self.writeSystemLog(u'登录失败!')
            raise Exception(u'登录失败!')

        # 执行时间变量的初始化
        runStartTime = datetime.now()
        currentTime = datetime.now()
        timeSpend = currentTime - runStartTime
        # 获取总出错项
        qs = RefreshRedo.objects.filter(state=u'redoing')
        redos = read_frame(qs)
        errorCount = len(redos)
        self.writeSystemLog(u'共有%d次错误需要重做' % errorCount)


        # 不停处理每一个信息项,直到没有错误需要重做,或者执行时间到了
        while errorCount > 0:
            try:
                qs = RefreshRedo.objects.filter(state=u'redoing')
                redos = read_frame(qs)
                errorCount = len(redos)
                if errorCount == 0:
                    self.writeSystemLog(u'已无项目需要处理,程序将退出')
                else:
                    self.writeSystemLog(u'开始一次重做，共%d需要处理...' % errorCount)

                # 根据出错列表调用重新回订过程
                successCount = 0
                # 按唯一的日期获取数据，提高访问效率
                for day in redos.beginDay.drop_duplicates():
                    self.writeSystemLog(u'检查%s的余票信息' % day)

                    # 获取有余票的项
                    ticketInfo = pitcher.getTicketInfo(day)
                    c1 = ticketInfo[u'余票'].apply(lambda x: int(x)) > 0
                    remain = ticketInfo[c1]
                    remain[u'开航日期'] = day

                    # 余票信息和重做记录关联，得出可以重做的票项
                    leftColumns = ['beginDay', 'beginTime', 'departure', 'arrival']
                    rightColumns = [u'开航日期', u'开航时间', u'出发码头', u'抵达码头']
                    todos = redos.merge(remain, left_on=leftColumns, right_on=rightColumns).to_dict(outtype='records')

                    if todos:
                        self.writeSystemLog(u'有%d个重做项出现余票' % len(todos))
                    else:
                        self.writeSystemLog(u'无重做项有余票')
                        # 避免过分频繁访问服务器
                        time.sleep(self.normalWaitingSecond)

                    # 循环对有余票的项目进程操作
                    for todo in todos:
                        # 初始化相关信息
                        departure = todo['departure']
                        arrival = todo['arrival']
                        beginDay = todo['beginDay']
                        beginTime = todo['beginTime']
                        cnt = todo['cnt']
                        remainCnt = int(todo[u'余票'])
                        redoId = todo['id']
                        redo = RefreshRedo.objects.get(id=redoId)
                        # 如果当前时间和开航日期0点的时间小于1天，则不再刷新这个票项
                        # 因为正常刷新提前5天前就刷新好了，所以不会有影响
                        # 这样做主要是防止重做票项无限增长，且去订之前日期的票
                        beginDateTime = parser.parse(beginDay)
                        if (beginDateTime - currentTime).days < 1:
                            redo.state = u'failed'
                            redo.save()
                            self.writeSystemLog(
                                u'出发:%s,抵达:%s,开航时间:%s %s,人数:%s' % (departure, arrival, beginDay, beginTime, cnt))
                            self.writeSystemLog(u'该项已经超过可以重做的时间，已标记为失败!!!')
                            continue

                        # 尝试重调订票过程
                        dailyFlightId = pitcher.getDailyFlightId(beginDay, beginTime, departure, arrival)
                        if dailyFlightId:
                            self.writeSystemLog(
                                u'出发:%s,抵达:%s,开航时间:%s %s,人数:%s' % (departure, arrival, beginDay, beginTime, cnt))
                            if remainCnt >= cnt:
                                # 如果票数量够，按需票量抢
                                result = pitcher.orderTicket(dailyFlightId, cnt)
                                if result:
                                    # 重做成功将记录标识为完成
                                    redo.state = u'finished'
                                    redo.save()
                                    self.writeSystemLog(u'该项已经重做成功!')
                                    successCount += 1
                                else:
                                    self.writeSystemLog(u'虽有余票项但抢订失败!')
                            else:
                                # 如果票数不够，按实际票数抢
                                result = pitcher.orderTicket(dailyFlightId, remainCnt)
                                if result:
                                    redo.cnt -= remainCnt
                                    redo.save()
                                    self.writeSystemLog(u'重做项余票数不够,但已抢回%d张' % remainCnt)
                                else:
                                    self.writeSystemLog(u'虽有余票项但抢订失败!')
                        # 避免过分频繁访问服务器(for todo1 in todos:)
                        time.sleep(self.normalWaitingSecond)
                # 避免过分频繁访问服务器(for day in redos.beginDay.drop_duplicates())
                time.sleep(self.normalWaitingSecond)

                # 完成一次重做,打印执行信息
                errorCount -= successCount
                self.writeSystemLog(u'完成一次重做：%d项成功，%d项失败.' % (successCount, errorCount))
                # 更新执行时间
                currentTime = datetime.now()
                timeSpend = currentTime - runStartTime
            except Exception as e:
                self.writeSystemLog(u'发生异常:%s，程序将继续执行' % unicode(e))

            # 检查是否是登录状态，如果登录状态丢失则重新登录
            if not self.isLogin():
                self.writeSystemLog(u'连接丢失，将重新登录...')
                loginErrorCount = 0
                while True:
                    loginResult = self.login(self.username, self.password)
                    if loginResult:
                        # 如果登录成功清空原来登录失败的记录
                        self.writeSystemLog(u'登录成功!')
                        break
                    else:
                        loginErrorCount += 1
                        if loginErrorCount > self.maxLoginError:
                            self.writeSystemLog(u'登录失败超过%d次，程序退出!' % self.maxLoginError)
                            return
                        else:
                            self.writeSystemLog(u'登录失败!')
                            self.writeSystemLog(u'等待%s秒... ...' % self.errorWaitingSecond)
                            time.sleep(self.errorWaitingSecond)

        self.writeSystemLog(u'重做过程执行结束')

    def run(self):
        '''
        更新主过程
        '''
        self.writeSystemLog(u'开始重做过程...')
        while True:
            try:
                self.runRedo()
            except Exception as e:
                self.writeSystemLog(u'发生异常:%s，程序将继续执行' % unicode(e))
            # 程序执行到这里有两种情况:
            # 1、抛出异常
            # 2、所有重做项执行完毕
            # 等待15分钟后重新开始执行
            time.sleep(900)
        self.writeSystemLog(u'重做过程已返回...')
